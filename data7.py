import streamlit as st
from PIL import Image, ImageDraw, ImageFont, ImageEnhance, ImageFilter
import PyPDF2
import os
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import textwrap
from io import BytesIO
import hashlib

st.title("Data Reader")
st.write("Upload a file to preview its contents")

BASE_DIR = "storage/files"
os.makedirs(BASE_DIR, exist_ok=True)

def pin_path(filename):
    return os.path.join(BASE_DIR, f"{filename}.pin")

def hash_pin(pin):
    return hashlib.sha256(pin.encode()).hexdigest()

def save_pin(filename, pin):
    with open(pin_path(filename), "w") as f:
        f.write(hash_pin(pin))

def verify_pin(filename, pin):
    try:
        with open(pin_path(filename), "r") as f:
            return f.read() == hash_pin(pin)
    except:
        return False

def require_pin(filename):
    pin = st.text_input("Enter PIN", type="password")
    if not pin or not verify_pin(filename, pin):
        st.error("Access denied")
        st.stop()

def file_path(name):
    return os.path.join(BASE_DIR, name)

def create_pdf_bytes(text):
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    y = height - 40
    for line in text.split("\n"):
        for w in textwrap.wrap(line, 90):
            if y < 40:
                c.showPage()
                y = height - 40
            c.drawString(40, y, w)
            y -= 15
    c.save()
    buf.seek(0)
    return buf

def get_files():
    return [f for f in os.listdir(BASE_DIR) if not f.endswith(".pin")]

def search_files(query):
    files = get_files()
    if query:
        files = [f for f in files if query.lower() in f.lower()]
    return files

st.divider()
st.header("Browse & Preview Files")

browse_file = st.file_uploader(
    "Upload a file to preview",
    type=["txt", "docx", "pdf", "jpg", "jpeg", "png"]
)

if browse_file:
    data = browse_file.read()
    if browse_file.name.endswith(".txt"):
        st.text(data.decode(errors="ignore")[:5000])
    elif browse_file.name.endswith(".docx"):
        doc = Document(BytesIO(data))
        st.text("\n".join(p.text for p in doc.paragraphs)[:5000])
    elif browse_file.name.endswith(".pdf"):
        reader = PyPDF2.PdfReader(BytesIO(data))
        st.text("".join(p.extract_text() or "" for p in reader.pages)[:5000])
    else:
        st.image(Image.open(BytesIO(data)), width=600)

st.divider()
st.header("File CRUD Operations")

crud_menu = st.radio(
    "Select Operation",
    ["Create (Save)", "Read", "Update", "Delete", "List Files"],
    horizontal=True
)

if crud_menu == "Create (Save)":
    f = st.file_uploader("Upload file")
    pin = st.text_input("Set PIN", type="password")
    if f and pin and st.button("Save"):
        with open(file_path(f.name), "wb") as out:
            out.write(f.getbuffer())
        save_pin(f.name, pin)
        st.success("File saved")

elif crud_menu == "Read":
    search = st.text_input("Search file")
    files = search_files(search)
    if files:
        name = st.selectbox("Select file", files)
        require_pin(name)
        path = file_path(name)

        if name.endswith(".txt"):
            st.text(open(path, errors="ignore").read()[:5000])
        elif name.endswith(".docx"):
            doc = Document(path)
            st.text("\n".join(p.text for p in doc.paragraphs)[:5000])
        elif name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(path)
            st.text("".join(p.extract_text() or "" for p in reader.pages)[:5000])
        else:
            st.image(Image.open(path), width=600)
        st.download_button("Download", open(path, "rb").read(), file_name=name)
    else:
        st.info("No files found")

elif crud_menu == "Update":
    search = st.text_input("Search file")
    files = search_files(search)
    if files:
        name = st.selectbox("Select file", files)
        require_pin(name)
        path = file_path(name)

        if name.endswith(".txt"):
            content = open(path, errors="ignore").read()
            edited = st.text_area("Edit TXT", content, height=400)
            if st.button("Update TXT"):
                open(path, "w", encoding="utf-8").write(edited)
                st.success("TXT updated")
            st.download_button("Download Updated TXT", edited, file_name=name)

        elif name.endswith(".docx"):
            doc = Document(path)
            content = "\n".join(p.text for p in doc.paragraphs)
            edited = st.text_area("Edit DOCX", content, height=400)
            if st.button("Update DOCX"):
                new_doc = Document()
                for line in edited.split("\n"):
                    new_doc.add_paragraph(line)
                new_doc.save(path)
                st.success("DOCX updated successfully")
            st.download_button("Download Updated DOCX", open(path, "rb").read(), file_name=name)

        elif name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(path)
            content = "".join(p.extract_text() or "" for p in reader.pages)
            edited = st.text_area("Edit PDF", content, height=400)
            if st.button("Generate Updated PDF"):
                pdf_buf = create_pdf_bytes(edited)
                st.download_button("Download Updated PDF", pdf_buf, file_name=name, mime="application/pdf")

        elif name.lower().endswith((".jpg", ".jpeg", ".png")):
            original_img = Image.open(path).convert("RGB")
            img = original_img.copy()
            st.image(img, caption="Original Image", width=600)

            st.subheader("Image Editing Options")


            brightness = st.slider("Brightness", 0.1, 2.0, 1.0)
            contrast = st.slider("Contrast", 0.1, 2.0, 1.0)
            sharpness = st.slider("Sharpness", 0.1, 2.0, 1.0)
            blur = st.slider("Blur", 0.0, 5.0, 0.0)

            text_layers = st.session_state.get("text_layers", [])

            new_text = st.text_input("Add New Text Layer")
            font_size = st.slider("Font Size", 10, 100, 30)
            font_color = st.color_picker("Font Color", "#FFFFFF")
            alignment = st.selectbox("Text Alignment", [
                "Top-Left", "Top-Center", "Top-Right",
                "Center-Left", "Center", "Center-Right",
                "Bottom-Left", "Bottom-Center", "Bottom-Right"
            ])

            def get_text_position(img, text, font, alignment, index):
                w, h = img.size
                try:
                    tw, th = font.getsize(text)
                except:
                    tw, th = len(text) * font_size // 2, font_size
                # Horizontal
                if "Left" in alignment:
                    x = 10
                elif "Center" in alignment:
                    x = (w - tw) // 2
                elif "Right" in alignment:
                    x = w - tw - 10
                # Vertical base
                if "Top" in alignment:
                    y = 10 + index * (th + 5)
                elif "Center" in alignment:
                    y = (h - th) // 2 + index * (th + 5)
                elif "Bottom" in alignment:
                    y = h - th - 10 - index * (th + 5)
                return x, y

            if st.button("Add Text Layer"):
                if new_text.strip():
                    text_layers.append({
                        "text": new_text.strip(),
                        "font_size": font_size,
                        "font_color": font_color,
                        "alignment": alignment
                    })
                    st.session_state["text_layers"] = text_layers
                    st.success("Text layer added!")

            if st.button("Preview Changes"):
                img = original_img.copy()
                img = ImageEnhance.Brightness(img).enhance(brightness)
                img = ImageEnhance.Contrast(img).enhance(contrast)
                img = ImageEnhance.Sharpness(img).enhance(sharpness)
                if blur > 0:
                    img = img.filter(ImageFilter.GaussianBlur(blur))

                draw = ImageDraw.Draw(img)
                for i, layer in enumerate(text_layers):
                    try:
                        font = ImageFont.truetype("arial.ttf", layer["font_size"])
                    except:
                        font = ImageFont.load_default()
                    pos = get_text_position(img, layer["text"], font, layer["alignment"], i)
                    draw.text(pos, layer["text"], font=font, fill=layer["font_color"])
                st.image(img, caption="Preview Edited Image", width=600)

            if st.button("Reset Image"):
                img = original_img.copy()
                st.session_state["text_layers"] = []
                st.image(img, caption="Reset to Original", width=600)

            # --- Download Edited Image ---
            buf = BytesIO()
            img.save(buf, format="PNG")
            buf.seek(0)
            st.download_button("Download Edited Image", buf, file_name=f"edited_{name}", mime="image/png")

elif crud_menu == "Delete":
    search = st.text_input("Search file")
    files = search_files(search)
    if files:
        name = st.selectbox("Select file to delete", files)
        require_pin(name)
        path = file_path(name)

        st.subheader("Preview before delete")
        if name.endswith(".txt"):
            st.text(open(path, errors="ignore").read()[:5000])
        elif name.endswith(".docx"):
            doc = Document(path)
            st.text("\n".join(p.text for p in doc.paragraphs)[:5000])
        elif name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(path)
            st.text("".join(p.extract_text() or "" for p in reader.pages)[:5000])
        else:
            st.image(Image.open(path), width=600)

        if st.button("Delete File"):
            os.remove(path)
            os.remove(pin_path(name))
            st.success("File deleted successfully")
    else:
        st.info("No files found")
elif crud_menu == "List Files":
    search = st.text_input("Search file")
    files = search_files(search)
    if files:
        name = st.selectbox("Select file", files)
        require_pin(name)
        path = file_path(name)

        st.subheader("File Preview")
        if name.endswith(".txt"):
            st.text(open(path, errors="ignore").read()[:5000])
        elif name.endswith(".docx"):
            doc = Document(path)
            st.text("\n".join(p.text for p in doc.paragraphs)[:5000])
        elif name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(path)
            st.text("".join(p.extract_text() or "" for p in reader.pages)[:5000])
        else:
            st.image(Image.open(path), width=600)

        st.download_button("Download File", open(path, "rb").read(), file_name=name)
    else:
        st.info("No files found")
        